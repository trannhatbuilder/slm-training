from pathlib import Path
from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent.parent
REPORT_PATH = PROJECT_ROOT / "data" / "raw" / "DOC-000001.docx"


def normalize_text(text: str) -> str:
    """Remove repeated whitespace without changing the source document."""
    return " ".join(text.split())


def is_heading(paragraph) -> bool:
    """Return True when a paragraph uses a Word heading style."""
    style_name = paragraph.style.name if paragraph.style else ""
    return style_name.lower().startswith("heading")


def main() -> None:
    if not REPORT_PATH.exists():
        raise FileNotFoundError(
            f"Source report was not found: {REPORT_PATH}"
        )

    document = Document(REPORT_PATH)

    non_empty_paragraphs = []
    headings = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = normalize_text(paragraph.text)

        if not text:
            continue

        non_empty_paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else None,
                "text": text,
            }
        )

        if is_heading(paragraph):
            headings.append(
                {
                    "index": index,
                    "style": paragraph.style.name,
                    "text": text,
                }
            )

    print("DOCX inspection completed")
    print(f"Report: {REPORT_PATH.name}")
    print(f"Total paragraphs: {len(document.paragraphs)}")
    print(f"Non-empty paragraphs: {len(non_empty_paragraphs)}")
    print(f"Total tables: {len(document.tables)}")
    print(f"Detected headings: {len(headings)}")

    print("\nDetected headings:")
    if not headings:
        print("- No Word heading styles were detected.")
    else:
        for heading in headings:
            print(
                f"- Paragraph {heading['index']} "
                f"[{heading['style']}]: {heading['text']}"
            )

    print("\nFirst 10 non-empty paragraphs:")
    for paragraph in non_empty_paragraphs[:10]:
        print(
            f"- Paragraph {paragraph['index']} "
            f"[{paragraph['style']}]: {paragraph['text']}"
        )


if __name__ == "__main__":
    main()