from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parent.parent
REPORT_PATH = PROJECT_ROOT / "data" / "raw" / "DOC-000001.docx"

FINDING_TITLES = [
    "Hardcoded RabbitMQ Credentials in Mobile Application",
    "Use of Symmetric JWT Signing Algorithm (HS256)",
    "Root Detection Not Implemented",
    "Weak SSL Pinning Implementation",
    "Email Enumeration Possible During User Registration",
]


def normalize_text(text: str) -> str:
    """Normalize whitespace only for inspection output."""
    return " ".join(text.split())


def text_matches_title(text: str, title: str) -> bool:
    """Perform a case-insensitive exact comparison after normalization."""
    return normalize_text(text).casefold() == normalize_text(title).casefold()


def inspect_paragraphs(document: Document) -> None:
    print("=" * 80)
    print("PARAGRAPH MATCHES")
    print("=" * 80)

    for title in FINDING_TITLES:
        matches = []

        for index, paragraph in enumerate(document.paragraphs, start=1):
            text = normalize_text(paragraph.text)

            if text_matches_title(text, title):
                matches.append(index)

        print(f"\nFinding title: {title}")

        if not matches:
            print("Paragraph match: NOT FOUND")
            continue

        for match_index in matches:
            paragraph = document.paragraphs[match_index - 1]
            style_name = (
                paragraph.style.name
                if paragraph.style
                else None
            )

            print(
                f"Paragraph match: {match_index} "
                f"[style={style_name}]"
            )

            start = max(1, match_index - 3)
            end = min(
                len(document.paragraphs),
                match_index + 8,
            )

            print(
                f"Context paragraphs: {start} to {end}"
            )

            for context_index in range(start, end + 1):
                context_paragraph = document.paragraphs[
                    context_index - 1
                ]
                context_text = normalize_text(
                    context_paragraph.text
                )

                if not context_text:
                    continue

                context_style = (
                    context_paragraph.style.name
                    if context_paragraph.style
                    else None
                )

                marker = (
                    ">>>"
                    if context_index == match_index
                    else "   "
                )

                print(
                    f"{marker} P{context_index} [{context_style}]: {context_text}"
                )

def inspect_tables(document: Document) -> None:
    print("\n" + "=" * 80)
    print("TABLE MATCHES")
    print("=" * 80)

    for title in FINDING_TITLES:
        matches = []

        for table_index, table in enumerate(
            document.tables,
            start=1,
        ):
            table_text_parts = []
            for row in table.rows:
                for cell in row.cells:
                    cell_text = normalize_text(cell.text)

                    if cell_text:
                        table_text_parts.append(cell_text)

            table_text = " | ".join(table_text_parts)

            if title.casefold() in table_text.casefold():
                matches.append(
                    {
                        "table_index": table_index,
                        "text": table_text,
                    }
                )

        print(f"\nFinding title: {title}")

        if not matches:
            print("Table match: NOT FOUND")
            continue

        for match in matches:
            print(
                f"Table match: {match['table_index']}"
            )
            print(
                f"Table preview: {match['text'][:1000]}"
            )


def list_custom_styles(document: Document) -> None:
    print("\n" + "=" * 80)
    print("NON-EMPTY PARAGRAPHS WITH EVVO STYLES")
    print("=" * 80)

    for index, paragraph in enumerate(
        document.paragraphs,
        start=1,
    ):
        text = normalize_text(paragraph.text)

        if not text:
            continue

        style_name = (
            paragraph.style.name
            if paragraph.style
            else ""
        )

        if style_name.casefold().startswith("evvo"):
            print(
                f"P{index} [{style_name}]: {text}"
            )


def main() -> None:
    if not REPORT_PATH.exists():
        raise FileNotFoundError(
            f"Source report was not found: {REPORT_PATH}"
        )

    document = Document(REPORT_PATH)

    print("Finding structure inspection")
    print(f"Report: {REPORT_PATH.name}")
    print(f"Paragraph count: {len(document.paragraphs)}")
    print(f"Table count: {len(document.tables)}")

    inspect_paragraphs(document)
    inspect_tables(document)
    list_custom_styles(document)


if __name__ == "__main__":
    main()